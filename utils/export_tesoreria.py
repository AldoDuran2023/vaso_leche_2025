import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from flask import send_file
from datetime import datetime

def export_to_pdf_tesoreria(data, filename="reporte_tesoreria.pdf"):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=70, bottomMargin=50)
    styles = getSampleStyleSheet()
    story = []

    # Título
    title = Paragraph(f"Reporte de Tesorería - Junta {data['anio_junta']}", styles['Title'])
    story.append(title)

    # Fecha
    fecha = Paragraph(f"Fecha del reporte: {datetime.now().strftime('%d/%m/%Y')}", styles['Normal'])
    story.append(fecha)
    story.append(Spacer(1, 12))

    def tabla_con_total(titulo, registros, total):
        story.append(Paragraph(f"{titulo}", styles['Heading2']))
        story.append(Spacer(1, 6))
        if not registros:
            story.append(Paragraph("No hay registros.", styles['Normal']))
            return

        columnas = list(registros[0].keys())
        data_table = [columnas] + [[r[col] for col in columnas] for r in registros]
        table = Table(data_table, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F9FA')]),
        ]))
        story.append(table)
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<b>Total: S/ {total:.2f}</b>", styles['Normal']))
        story.append(Spacer(1, 20))

    # Ingresos
    tabla_con_total("Ingresos", data['ingresos'], data['total_ingresos'])

    # Gastos
    tabla_con_total("Gastos", data['gastos'], data['total_gastos'])

    doc.build(story)
    buffer.seek(0)

    return send_file(
        buffer,
        download_name=filename,
        as_attachment=True,
        mimetype='application/pdf'
    )

def export_to_excel_tesoreria(data, filename="reporte_tesoreria.xlsx"):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # --- Ingresos ---
        ingresos_df = pd.DataFrame(data['ingresos'])
        if not ingresos_df.empty:
            ingresos_df.to_excel(writer, sheet_name='Ingresos', index=False, startrow=4)
            ws = writer.sheets['Ingresos']
            ws['A1'] = f"Reporte de Tesorería - Junta {data['anio_junta']}"
            ws['A2'] = f"Fecha del reporte: {datetime.now().strftime('%d/%m/%Y')}"
            ws['A4'] = "INGRESOS"
            ws['A4'].font = ws['A4'].font.copy(bold=True)

            total_row = len(ingresos_df) + 6
            ws[f"A{total_row}"] = "Total"
            ws[f"D{total_row}"] = f"S/ {data['total_ingresos']:.2f}"

        # --- Gastos ---
        gastos_df = pd.DataFrame(data['gastos'])
        if not gastos_df.empty:
            gastos_df.to_excel(writer, sheet_name='Gastos', index=False, startrow=4)
            ws2 = writer.sheets['Gastos']
            ws2['A1'] = f"Reporte de Tesorería - Junta {data['anio_junta']}"
            ws2['A2'] = f"Fecha del reporte: {datetime.now().strftime('%d/%m/%Y')}"
            ws2['A4'] = "EGRESOS"
            ws2['A4'].font = ws2['A4'].font.copy(bold=True)

            total_row = len(gastos_df) + 6
            ws2[f"A{total_row}"] = "Total"
            ws2[f"D{total_row}"] = f"S/ {data['total_gastos']:.2f}"

    output.seek(0)
    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

def export_to_word_tesoreria(data, filename="reporte_tesoreria.docx"):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Título y fecha
    title = doc.add_heading(f"Reporte de Tesorería - Junta {data['anio_junta']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha = doc.add_paragraph(f"Fecha del reporte: {datetime.now().strftime('%d/%m/%Y')}")
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("")

    def tabla_con_total(titulo, registros, total):
        doc.add_heading(titulo, level=1)
        if not registros:
            doc.add_paragraph("No hay registros.")
            return

        columnas = list(registros[0].keys())
        table = doc.add_table(rows=1, cols=len(columnas))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Encabezados
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columnas):
            hdr_cells[i].text = str(col)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(10)

        # Filas
        for row in registros:
            cells = table.add_row().cells
            for i, col in enumerate(columnas):
                text = str(row[col])
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.size = Pt(9)

        doc.add_paragraph(f"\nTotal: S/ {total:.2f}", style='Normal')

    tabla_con_total("Ingresos", data['ingresos'], data['total_ingresos'])
    doc.add_paragraph("\n")
    tabla_con_total("Gastos", data['gastos'], data['total_gastos'])

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
