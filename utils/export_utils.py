import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from flask import send_file
from datetime import datetime

def export_to_excel(data: list[dict], title: str = "Reporte", filename: str = "reporte.xlsx"):
    """Exportar a Excel con formato profesional"""
    if not data:
        return None
    
    df = pd.DataFrame(data)
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Escribir datos
        df.to_excel(writer, sheet_name='Reporte', index=False, startrow=3)
        
        # Obtener la hoja de trabajo
        worksheet = writer.sheets['Reporte']
        
        # Agregar título y fecha
        worksheet['A1'] = title
        worksheet['A2'] = f"Fecha del reporte: {datetime.now().strftime('%d/%m/%Y')}"
        
        # Formatear título
        title_cell = worksheet['A1']
        title_cell.font = title_cell.font.copy(bold=True, size=14)
        
        # Formatear encabezados
        for col_num, column_title in enumerate(df.columns, 1):
            cell = worksheet.cell(row=4, column=col_num)
            cell.font = cell.font.copy(bold=True)
            cell.fill = cell.fill.copy(fgColor="D9D9D9")
        
        # Ajustar ancho de columnas automáticamente
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            adjusted_width = min(max_length + 2, 50)  # Máximo 50 caracteres
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    output.seek(0)
    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

def export_to_word(data: list[dict], title: str = "Reporte", filename: str = "reporte.docx"):
    """Exportar a Word con formato profesional"""
    if not data:
        return None
    
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Título principal
    title_paragraph = doc.add_heading(title, 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha
    fecha_paragraph = doc.add_paragraph(f"Fecha del reporte: {datetime.now().strftime('%d/%m/%Y')}")
    fecha_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Espacio
    doc.add_paragraph("")
    
    columns = list(data[0].keys())
    
    # Crear tabla
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Configurar encabezados
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(columns):
        hdr_cells[i].text = str(column).title()
        # Formatear encabezado
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(10)
    
    # Agregar datos
    for row_data in data:
        row_cells = table.add_row().cells
        for i, column in enumerate(columns):
            cell_text = str(row_data.get(column, ''))
            row_cells[i].text = cell_text
            # Centrar contenido
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Ajustar tamaño de fuente
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    
    # Ajustar ancho de columnas proporcionalmente
    total_width = Inches(7)  # Ancho total disponible
    col_width = total_width / len(columns)
    
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width
    
    # Ajustar espaciado de filas
    for row in table.rows:
        row.height = Inches(0.3)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def export_to_pdf(data: list[dict], title: str = "Reporte", filename: str = "reporte.pdf"):
    """Exportar a PDF con formato profesional y paginación automática"""
    if not data:
        return None
    
    buffer = BytesIO()
    
    # Usar A4 para mejor formato
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        rightMargin=50, 
        leftMargin=50,
        topMargin=70, 
        bottomMargin=50
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Estilo personalizado para el título
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        spaceAfter=12,
        alignment=1,  # Centrado
        textColor=colors.black
    )
    
    # Estilo para la fecha
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=10,
        alignment=2,  # Derecha
        spaceAfter=20
    )
    
    # Título
    story.append(Paragraph(title, title_style))
    
    # Fecha
    fecha = datetime.now().strftime("%d/%m/%Y")
    story.append(Paragraph(f"Fecha del reporte: {fecha}", date_style))
    
    # Preparar datos para la tabla
    columns = list(data[0].keys())
    
    # Ajustar ancho de página para A4
    page_width = A4[0] - 100  # Restamos márgenes
    
    # Calcular ancho de columnas dinámicamente
    num_cols = len(columns)
    
    # Ancho mínimo y máximo por columna
    min_col_width = 60
    max_col_width = page_width / 3  # Máximo 3 columnas por ancho de página
    
    # Calcular ancho óptimo
    optimal_width = page_width / num_cols
    col_width = max(min_col_width, min(optimal_width, max_col_width))
    
    # Si las columnas son muy anchas, usar el ancho de página completo
    if col_width * num_cols > page_width:
        col_width = page_width / num_cols
    
    col_widths = [col_width] * num_cols
    
    # Preparar encabezados con formato
    headers = [str(col).title() for col in columns]
    table_data = [headers]
    
    # Agregar datos con formato adecuado
    for row in data:
        formatted_row = []
        for col in columns:
            value = row.get(col, '')
            # Limitar longitud de texto para evitar desbordamiento
            formatted_value = str(value)
            if len(formatted_value) > 25:  # Truncar texto muy largo
                formatted_value = formatted_value[:22] + "..."
            formatted_row.append(formatted_value)
        table_data.append(formatted_row)
    
    # Crear tabla con paginación automática
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    
    # Estilo profesional de la tabla
    table_style = TableStyle([
        # Encabezados
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        
        # Contenido
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        
        # Filas alternadas
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F9FA')]),
    ])
    
    table.setStyle(table_style)
    story.append(table)
    
    # Construir documento
    doc.build(story)
    buffer.seek(0)
    
    return send_file(
        buffer,
        download_name=filename,
        as_attachment=True,
        mimetype='application/pdf'
    )

